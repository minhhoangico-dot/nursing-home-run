"""
One-off generator: convert the lawyer-approved contract sample
(`docs/HĐ MẪU 3 PHỤ LỤC 4.2026.docx`) into two docxtemplater-ready files:

  - `public/templates/contract_template_v1.docx` — full contract
  - `public/templates/assessment_template_v1.docx` — Phụ lục 2 only
    (Phiếu đánh giá mức độ chăm sóc NCT)

Run from repo root:
    python scripts/build-contract-template.py

Re-run whenever the source contract changes.
"""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

SOURCE = Path("docs/HĐ MẪU 3 PHỤ LỤC 4.2026.docx")
DEST = Path("public/templates/contract_template_v1.docx")
DEST_ASSESSMENT = Path("public/templates/assessment_template_v1.docx")


def replace_in_runs(paragraph, old: str, new: str) -> bool:
    """
    Replace `old` with `new` across consecutive runs, preserving formatting
    outside the matched span. The replacement inherits the formatting of
    the first matched run.
    """
    runs = list(paragraph.runs)
    if not runs:
        return False

    cum = 0
    starts = []
    for r in runs:
        starts.append(cum)
        cum += len(r.text)
    text = "".join(r.text for r in runs)

    idx = text.find(old)
    if idx == -1:
        return False
    end = idx + len(old)

    first = next(
        (i for i, s in enumerate(starts) if s <= idx < s + len(runs[i].text)),
        None,
    )
    last = next(
        (i for i, s in enumerate(starts) if s < end <= s + len(runs[i].text)),
        None,
    )
    if first is None or last is None:
        return False

    pre = runs[first].text[: idx - starts[first]]
    post = runs[last].text[end - starts[last] :]
    runs[first].text = pre + new + post
    for i in range(first + 1, last + 1):
        runs[i]._element.getparent().remove(runs[i]._element)
    return True


def set_cell_text(cell, text: str) -> None:
    """Replace cell content with a single run carrying `text` (preserves cell paragraph style)."""
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)
    # Drop extra paragraphs if any
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def main() -> None:
    if not SOURCE.exists():
        raise SystemExit(f"Source template not found: {SOURCE}")
    DEST.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(SOURCE))

    # ── Header table (Table 0) ──────────────────────────────────────────
    t0 = doc.tables[0]
    # R0C0 P1 "Số:     /2026/DV-VDL" -> placeholder
    p_num = t0.rows[0].cells[0].paragraphs[1]
    replace_in_runs(p_num, "Số:     /2026/DV-VDL", "Số: {contract_number}/2026/DV-VDL")
    # R0C1 P2 "Hà Nội, ngày    tháng    năm2026"
    p_date = t0.rows[0].cells[1].paragraphs[2]
    replace_in_runs(
        p_date,
        "Hà Nội, ngày    tháng    năm2026",
        "Hà Nội, ngày {signed_day} tháng {signed_month} năm {signed_year}",
    )

    # ── Body paragraph: ngày ký (P7) ────────────────────────────────────
    # "...vào ngày   tháng      năm 2026,..."
    p7 = doc.paragraphs[7]
    replace_in_runs(
        p7,
        "ngày   tháng      năm 2026",
        "ngày {signed_day} tháng {signed_month} năm {signed_year}",
    )

    # ── Điều 2 (P22): tên NCT ──────────────────────────────────────────
    replace_in_runs(doc.paragraphs[22], "Nhữ Thị Kim Dung", "{resident_name}")

    # ── Điều 2 (P23): địa chỉ NCT ──────────────────────────────────────
    replace_in_runs(
        doc.paragraphs[23],
        "Địa chỉ: Thôn Long Phú, Xã Phú Cát, TP. Hà Nội, VN",
        "Địa chỉ: {resident_address}",
    )

    # ── Bảng người bảo trợ (Table 1) ───────────────────────────────────
    t1 = doc.tables[1]
    set_cell_text(t1.rows[0].cells[1], ": {guardian_name}")
    set_cell_text(t1.rows[1].cells[1], ": {guardian_dob}")
    set_cell_text(t1.rows[2].cells[1], ": {guardian_address}")
    set_cell_text(t1.rows[3].cells[1], ": {guardian_phone}")
    set_cell_text(t1.rows[4].cells[1], ": {guardian_id_card}")
    set_cell_text(t1.rows[5].cells[1], ": {guardian_relation}")

    # ── Bảng người cao tuổi (Table 2) ──────────────────────────────────
    t2 = doc.tables[2]
    set_cell_text(t2.rows[0].cells[1], ": {resident_name}")
    set_cell_text(t2.rows[1].cells[1], ": {resident_dob}")
    set_cell_text(t2.rows[2].cells[1], ": {resident_address}")
    set_cell_text(t2.rows[3].cells[1], ": {resident_phone}")
    set_cell_text(t2.rows[4].cells[1], ": {resident_id_card}")

    # ── Phụ lục headers (P126, P187, P236) ────────────────────────────
    p126 = doc.paragraphs[126]
    replace_in_runs(
        p126,
        "(Kèm theo thỏa thuận số          DV-VDL ngày         )",
        "(Kèm theo thỏa thuận số {contract_number}/2026/DV-VDL ngày {signed_date_full})",
    )

    p187 = doc.paragraphs[187]
    replace_in_runs(
        p187,
        "(Kèm theo thỏa thuận số         /202 /DV-VDL ngày)",
        "(Kèm theo thỏa thuận số {contract_number}/2026/DV-VDL ngày {signed_date_full})",
    )

    p236 = doc.paragraphs[236]
    replace_in_runs(
        p236,
        "(Kèm theo thỏa thuận số /2026 /DV-VDL ngày)",
        "(Kèm theo thỏa thuận số {contract_number}/2026/DV-VDL ngày {signed_date_full})",
    )

    # ── Phụ lục 3 — tên NCT trong cam kết ──────────────────────────────
    # P238: "...gửi người thân vào để hỗ trợ chăm sóc người nhà tôi là ông/bà: ........"
    for p in doc.paragraphs:
        if "gửi người thân vào để hỗ trợ chăm sóc người nhà tôi là ông/bà" in p.text:
            # Replace the trailing dotted blank with a placeholder
            replace_in_runs(
                p,
                "..................................... ………………………………………….",
                "{resident_name}",
            )
            break

    doc.save(str(DEST))
    print(f"Wrote {DEST}")

    build_assessment_template()


# ── Phụ lục 2: Phiếu đánh giá cấp độ chăm sóc ─────────────────────────────
# Spans body children from the paragraph "ĐÁNH GIÁ MỨC ĐỘ CHĂM SÓC NGƯỜI CAO TUỔI"
# through the signature table that immediately follows P232. Detected by
# scanning paragraph text for the start/end markers so that small edits to
# the contract source (extra blank paragraphs, etc.) don't break extraction.

ASSESSMENT_START_TEXT = "ĐÁNH GIÁ MỨC ĐỘ CHĂM SÓC NGƯỜI CAO TUỔI"
ASSESSMENT_NEXT_PHU_LUC_TEXT = "CAM KẾT CỦA BÊN A"


def build_assessment_template() -> None:
    doc = Document(str(SOURCE))
    body = doc.element.body

    children = list(body.iterchildren())

    start_idx = None
    end_idx = None  # exclusive — first child that belongs to Phụ lục 3
    for idx, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.iter(qn("w:t")))
        if start_idx is None and ASSESSMENT_START_TEXT in text:
            start_idx = idx
        elif start_idx is not None and ASSESSMENT_NEXT_PHU_LUC_TEXT in text:
            end_idx = idx
            break

    if start_idx is None or end_idx is None:
        raise SystemExit(
            f"Could not locate Phụ lục 2 boundaries in {SOURCE}. "
            f"start={start_idx}, end={end_idx}"
        )

    # Trim trailing empty paragraphs and the signature table that belongs
    # to the next section's page break. We keep everything up to (but not
    # including) end_idx, then walk back to drop trailing blanks.
    keep_until = end_idx - 1
    while keep_until > start_idx:
        c = children[keep_until]
        if c.tag == qn("w:p"):
            text = "".join(t.text or "" for t in c.iter(qn("w:t"))).strip()
            if not text:
                keep_until -= 1
                continue
        break

    # Remove children outside [start_idx, keep_until]
    for idx, child in enumerate(children):
        if idx < start_idx or idx > keep_until:
            parent = child.getparent()
            if parent is not None:
                parent.remove(child)

    # Apply placeholder substitutions on the remaining content.
    # Header "(Kèm theo thỏa thuận số ...)" becomes contract-aware.
    for p in doc.paragraphs:
        if p.text.startswith("(Kèm theo thỏa thuận số"):
            replace_in_runs(
                p,
                "(Kèm theo thỏa thuận số         /202 /DV-VDL ngày)",
                "(Kèm theo thỏa thuận số {contract_number}/2026/DV-VDL ngày {signed_date_full})",
            )
            break

    # Resident metadata fields (P190-P193 in source). Match by line prefix
    # because dotted fillers vary slightly.
    field_replacements = [
        ("Họ và tên người NCT:", "Họ và tên người NCT: {resident_name}"),
        ("Ngày tháng năm sinh:", "Ngày tháng năm sinh: {resident_dob}"),
        ("Ngày đánh giá:", "Ngày đánh giá: {assessment_date}"),
        ("Người đánh giá (Họ tên/Chức vụ):", "Người đánh giá (Họ tên/Chức vụ): {assessor_name}"),
    ]

    for p in doc.paragraphs:
        text = p.text
        for prefix, replacement in field_replacements:
            if text.startswith(prefix):
                # Replace whole paragraph text via runs
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
                p.add_run(replacement)
                break

    DEST_ASSESSMENT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DEST_ASSESSMENT))
    print(f"Wrote {DEST_ASSESSMENT}")


if __name__ == "__main__":
    main()
